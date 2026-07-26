#!/usr/bin/env python3
"""Generate a professional, editable DOCX resume."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# --- Page setup ---
section = doc.sections[0]
section.top_margin = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# --- Helper functions ---
def add_section_heading(text):
    p = doc.add_paragraph()
    p.space_before = Pt(18)
    p.space_after = Pt(8)
    run = p.add_run(text)
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1A, 0x3A, 0x5C)
    run.font.bold = True
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    # Add bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), '2C5F8A')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def add_normal_text(text, bold=False, size=10.5, color=None, indent=False):
    p = doc.add_paragraph()
    p.space_before = Pt(2)
    p.space_after = Pt(2)
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if bold:
        run.font.bold = True
    if color:
        run.font.color.rgb = color
    return p

def add_bullet(text, level=0, bold_prefix=None):
    p = doc.add_paragraph()
    p.space_before = Pt(1)
    p.space_after = Pt(1)
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    p.style = doc.styles['List Bullet']
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.font.bold = True
        run_b.font.size = Pt(10.5)
        run_b.font.name = '微软雅黑'
        run_b._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return p

def add_highlight_text(text):
    p = doc.add_paragraph()
    p.space_before = Pt(2)
    p.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
    run.font.bold = True
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    return p

# ========== HEADER ==========
name_p = doc.add_paragraph()
name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
name_p.space_after = Pt(2)
run = name_p.add_run('[姓名]')
run.font.size = Pt(26)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x3A, 0x5C)
run.font.name = '微软雅黑'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_p.space_after = Pt(6)
run = sub_p.add_run('智能制造 / 数字化转型 ｜ 资深总监')
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
run.font.name = '微软雅黑'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# Contact info
contact_p = doc.add_paragraph()
contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
contact_p.space_after = Pt(4)
run = contact_p.add_run('📧 [邮箱]　｜　📱 [手机]　｜　📍 上海　｜　🎂 [年龄]')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x77, 0x77, 0x77)
run.font.name = '微软雅黑'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# Divider
div_p = doc.add_paragraph()
div_p.space_before = Pt(2)
div_p.space_after = Pt(2)
pPr = div_p._p.get_or_add_pPr()
pBdr = OxmlElement('w:pBdr')
bottom = OxmlElement('w:bottom')
bottom.set(qn('w:val'), 'single')
bottom.set(qn('w:sz'), '6')
bottom.set(qn('w:space'), '1')
bottom.set(qn('w:color'), '1A3A5C')
pBdr.append(bottom)
pPr.append(pBdr)

# ========== 个人概述 ==========
add_section_heading('个人概述')
add_normal_text(
    '17+年制造业经验，其中6年智能制造数字化总监，兼具传统汽车制造（比亚迪2.5年+上汽通用9年）'
    '与智能制造跨界背景。上海交通大学机械设计及自动化硕士，智能制造方向博士在读。'
    '擅长从0到1推动数字化工厂建设，主导过数字化工艺全链路落地项目，带领60人跨职能团队'
    '（销售+售前+技术+项目），具备技术深度、市场开拓与团队管理的复合能力。'
    '英语可作为工作语言，拥有海外工作经历。'
)

# ========== 核心能力 ==========
add_section_heading('核心能力')
skills = ['智能制造', '数字化转型', 'MES', 'IoT', 'AI应用', '数字化工艺', '项目管理',
          '团队管理(60人)', '大客户拓展', '汽车制造工艺', '商务谈判', '英语工作语言', '产线规划']
skill_text = '、'.join(skills)
add_normal_text(skill_text, indent=True)

# ========== 工作经历 ==========
add_section_heading('工作经历')

def add_job_entry(role, company, period, bullets, highlights=None):
    # Role + company line
    p = doc.add_paragraph()
    p.space_before = Pt(10)
    p.space_after = Pt(2)
    run_role = p.add_run(role)
    run_role.font.bold = True
    run_role.font.size = Pt(11)
    run_role.font.name = '微软雅黑'
    run_role._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run_sep = p.add_run('　｜　')
    run_sep.font.size = Pt(10)
    run_company = p.add_run(company)
    run_company.font.size = Pt(10)
    run_company.font.color.rgb = RGBColor(0x2C, 0x5F, 0x8A)
    run_company.font.name = '微软雅黑'
    run_company._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run_date = p.add_run(f'　{period}')
    run_date.font.size = Pt(10)
    run_date.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run_date.font.name = '微软雅黑'
    run_date._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    for b in bullets:
        add_bullet(b)

# Current job
add_job_entry(
    '数字化总监 / 市场总监',
    '[当前公司] · 智能制造行业',
    '2020.07 – 至今（6年）',
    [
        '主导从3D图纸标注→工艺规划→程序生成→程序下发→试切的端到端数字化工艺平台建设，实现设计与制造一体化',
        '通过临港政府专项考核，获得1600万元现金补助',
        '自主完成潍柴动力及中国航发黎阳各2亿元项目的商务攻关与落地，累计合同金额超4亿',
        '带领公司在机加工基础上，成功拓展装配、焊接、数字化等新行业方向',
        '管理60人跨职能团队（销售、售前方案、技术研发、项目管理），直接下属5人',
        '熟悉MES、IoT平台架构，具备利用AI工具快速搭建系统Demo的能力',
    ]
)

add_job_entry(
    '工艺工程师 / 项目管理工程师',
    '上海通用汽车有限公司',
    '[约2011] – [约2020]（9年）',
    [
        '负责整车制造工艺规划与优化，涵盖机加工、装配等核心工艺环节',
        '担任多个新车型导入项目的项目管理，协调跨部门资源确保项目按时交付',
        '深入理解汽车行业精益生产和质量管理体系',
    ]
)

add_job_entry(
    '工艺工程师',
    '比亚迪股份有限公司',
    '[约2008] – [约2011]（2.5年）',
    [
        '从事汽车制造工艺相关工作，积累了扎实的现场工艺和制造工程经验',
    ]
)

# ========== 重点项目 ==========
add_section_heading('重点项目')

def add_project(title, desc, highlight=None):
    p = doc.add_paragraph()
    p.space_before = Pt(6)
    p.space_after = Pt(1)
    run = p.add_run(title)
    run.font.bold = True
    run.font.size = Pt(11)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    add_normal_text(desc, indent=True)
    if highlight:
        add_highlight_text(highlight)

add_project(
    '🏭 数字化工艺全链路平台',
    '主导建设覆盖"3D标注→工艺规划→程序生成→程序下发→试切验证"的端到端数字化工艺平台，打通设计与制造数据链，实现工艺知识沉淀与复用。项目通过临港政府专项验收，成为区域智能制造标杆案例。',
    '★ 获得临港政府1600万元现金补助'
)

add_project(
    '🤝 潍柴动力 · 2亿元项目',
    '独立主导潍柴动力智能制造项目的商务攻关、方案设计和项目交付全流程，合同金额2亿元。深度理解大型国企客户需求与决策链，成功建立长期合作关系。'
)

add_project(
    '✈️ 中国航发黎阳 · 2亿元项目',
    '独立完成中国航发黎阳项目的市场拓展与项目落地，合同金额2亿元。将公司在机加工领域的核心能力成功导入航空航天高端制造领域。'
)

add_project(
    '📈 行业多元化拓展',
    '统筹规划并推动公司业务从单一机加工向装配、焊接、数字化等多元行业方向延伸，拓宽市场覆盖与技术能力边界。'
)

# ========== 教育背景 ==========
add_section_heading('教育背景')

edu_data = [
    ('博士在读 · 智能制造方向', '[学校名称]', '预计2028年毕业'),
    ('硕士 · 机械设计及其自动化', '上海交通大学', '[毕业年份]'),
    ('本科 · 机械设计及其自动化', '[本科院校]', '[毕业年份]'),
]

for degree, school, year in edu_data:
    p = doc.add_paragraph()
    p.space_before = Pt(3)
    p.space_after = Pt(1)
    run_d = p.add_run(degree)
    run_d.font.bold = True
    run_d.font.size = Pt(10.5)
    run_d.font.name = '微软雅黑'
    run_d._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    run_s = p.add_run(f'　｜　{school}　｜　{year}')
    run_s.font.size = Pt(10)
    run_s.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    run_s.font.name = '微软雅黑'
    run_s._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# ========== 荣誉与资质 ==========
add_section_heading('荣誉与资质')

honors = [
    '🏆 汽车行业科技进步一等奖',
    '📜 中级工程师认证',
    '🌍 英语：商务谈判 & 技术交流（工作语言），拥有海外工作经历',
]

for h in honors:
    add_bullet(h)

# ========== 求职意向 ==========
add_section_heading('求职意向')
add_normal_text('意向岗位：数字化总监 / 智能制造总监 / 数字化转型负责人', indent=True)
add_normal_text('期望地点：上海', indent=True)
add_normal_text('期望行业：智能制造、汽车、航空航天、新能源、半导体等制造业（开放）', indent=True)
add_normal_text('公司类型：大型国企 > 大型民企 > 外资企业', indent=True)

# ========== Footer note ==========
p = doc.add_paragraph()
p.space_before = Pt(20)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('* 本简历为初稿，请根据实际情况补充 [ ] 标注的信息后使用')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)
run.font.name = '微软雅黑'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

# Save
output_path = os.path.join(os.path.dirname(__file__), '简历初稿.docx')
doc.save(output_path)
print(f'✅ DOCX简历已生成: {output_path}')
